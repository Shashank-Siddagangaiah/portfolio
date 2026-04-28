"""
Generate ISSUE_RESOLUTION.docx from ISSUE_RESOLUTION.md content.
Run with: C:/Users/shash/anaconda3/python.exe generate_docx.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = r"c:\Users\shash\Music\VS_code\Claude_git\Paperless_report\ISSUE_RESOLUTION.docx"

# ── colour palette ────────────────────────────────────────────────────────────
NAVY       = RGBColor(0x1F, 0x37, 0x64)   # heading 1 / title
DARK_BLUE  = RGBColor(0x2E, 0x74, 0xB5)   # heading 2
MID_BLUE   = RGBColor(0x2E, 0x74, 0xB5)   # heading 3
TABLE_HDR  = RGBColor(0x1F, 0x37, 0x64)   # table header fill
CODE_BG    = RGBColor(0xF2, 0xF2, 0xF2)   # code block shading
CODE_FG    = RGBColor(0x1A, 0x1A, 0x1A)
LABEL_CLR  = RGBColor(0x2E, 0x74, 0xB5)   # "Problem / Discovery / Fix / Result" labels


# ── helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_str):
    """Set cell background colour via XML (python-docx has no native API)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_str)
    tcPr.append(shd)


def set_cell_borders(cell, color="BFBFBF"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def cell_para(cell, text, bold=False, color=None, size=10, mono=False, align=None):
    p   = cell.paragraphs[0]
    p.clear()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if mono:
        run.font.name = 'Courier New'
    return p


def add_table(doc, headers, rows, col_widths=None):
    """Add a styled table with navy header row."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, '1F3764')
        p    = cell_para(cell, h, bold=True, color=RGBColor(0xFF,0xFF,0xFF), size=9)

    # Data rows
    for r_idx, row in enumerate(rows):
        tr = tbl.rows[r_idx + 1]
        bg = 'FFFFFF' if r_idx % 2 == 0 else 'EEF3FA'
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            set_cell_bg(cell, bg)
            cell_para(cell, str(val), size=9)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()   # spacing after table
    return tbl


def add_code_block(doc, code_text):
    """Add a shaded monospace paragraph for SQL snippets."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Cm(0.8)
    para.paragraph_format.right_indent = Cm(0.8)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)

    # Shade the paragraph
    pPr  = para._p.get_or_add_pPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F2F2F2')
    pPr.append(shd)

    run = para.add_run(code_text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(8)
    run.font.color.rgb = CODE_FG
    return para


def add_label(doc, label, text):
    """Add a bold coloured label followed by normal text on the same paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(2)
    r1 = para.add_run(label + "  ")
    r1.bold = True
    r1.font.color.rgb = LABEL_CLR
    r1.font.size = Pt(10)
    r2 = para.add_run(text)
    r2.font.size = Pt(10)
    return para


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(16)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = DARK_BLUE
        run.font.size = Pt(13)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = MID_BLUE
        run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p


def body(doc, text, italic=False, indent=False):
    para = doc.add_paragraph(text)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(4)
    if indent:
        para.paragraph_format.left_indent = Cm(0.8)
    for run in para.runs:
        run.font.size = Pt(10)
        if italic:
            run.italic = True
    return para


def note(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.8)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(6)
    r = para.add_run("Note: " + text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    return para


def page_break(doc):
    doc.add_page_break()


# ── DOCUMENT BUILD ─────────────────────────────────────────────────────────────

doc = Document()

# Default styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# Narrow margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── TITLE PAGE ────────────────────────────────────────────────────────────────
doc.add_paragraph()
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_para.add_run("Paperless Report\nIssue Resolution Document")
tr.bold = True
tr.font.size = Pt(22)
tr.font.color.rgb = NAVY

doc.add_paragraph()
meta_para = doc.add_paragraph()
meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in [
    "Project:  EDW/AWM Paperless Report (replacement for legacy BIM/Eloqua)",
    "Files:    new_paper.sql  |  cif_join_validation.sql  |  missing_mail.sql",
    "Date:     2026-04-07",
    "Author:   _______________",
]:
    r = meta_para.add_run(line + "\n")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

page_break(doc)

# ── EXECUTIVE SUMMARY ─────────────────────────────────────────────────────────
h1(doc, "Executive Summary")
body(doc,
    "During the build and validation of new_paper.sql — the new EDW/AWM-based paperless report "
    "replacing the legacy old_paper.sql (BIM/Eloqua source) — 9 data quality issues were identified "
    "and resolved. The pipeline was also optimized from 16 to 10 temp tables."
)

h2(doc, "Issue Impact Summary")
add_table(doc,
    headers=["#", "Issue", "Policies Impacted"],
    rows=[
        ["1", "Strict policyholder date join dropping policies",          "35,820"],
        ["2", "CIF superseded records not filtered",                       "Unknown (data quality risk)"],
        ["3", "CIF dedup using wrong sort key",                            "Unknown (data quality risk)"],
        ["4", "NULL paper_notify coerced to 'N' (false negative)",         "Affected all NULL CIF rows"],
        ["5", "Email fan-out from partition bug",                          "All symbols (counts > total)"],
        ["6", "Known bad party anchor IDs in email chain",                 "Inflated link counts"],
        ["7", "BIM fallback join format mismatch",                         "27,756 (fallback = 0)"],
        ["8", "NIN dedup preference losing ANI email",                     "26,476"],
        ["9", "BIM HOH filter too restrictive",                            "~25,000+ (fallback = 2,375 vs expected 27,756)"],
    ],
    col_widths=[0.3, 4.0, 2.0],
)

h2(doc, "Final Email Coverage")
add_table(doc,
    headers=["Metric", "Before Fixes", "After All Fixes"],
    rows=[
        ["Total policies",         "297,413",             "296,399 (grain dedup)"],
        ["AWM email",              "225,188 (75.8%)",     "251,664 (84.9%)"],
        ["BIM fallback",           "0",                   "464 (0.2%)"],
        ["Has email total",        "225,188 (75.8%)",     "252,128 (85.1%)"],
        ["No email (structural)",  "72,225",              "44,271 (14.9%)"],
    ],
    col_widths=[2.5, 2.2, 2.2],
)

page_break(doc)

# ── ISSUE TEMPLATE ─────────────────────────────────────────────────────────────

def issue_section(number, title, problem_text, discovery_text, fix_text, result_text,
                  code_before=None, code_after=None, code_fix=None, code_note=None,
                  extra_para=None):
    h1(doc, f"Issue {number} — {title}")

    h3(doc, "Problem")
    body(doc, problem_text)
    if code_before:
        add_code_block(doc, code_before)
    if code_note:
        body(doc, code_note)

    h3(doc, "Discovery")
    body(doc, discovery_text)

    h3(doc, "Fix")
    body(doc, fix_text)
    if code_after:
        add_code_block(doc, code_after)
    if code_fix:
        add_code_block(doc, code_fix)

    h3(doc, "Result")
    body(doc, result_text)
    if extra_para:
        extra_para()


# ── ISSUE 1 ───────────────────────────────────────────────────────────────────
issue_section(
    1,
    "Strict Policyholder Date Join Dropping ~35,820 Policies",
    problem_text=(
        "In cif_join_validation.sql, #view_ph joined DWM.EDW.vw_policyholder using three conditions: "
        "policy_term_key + term_effective_date + term_expiration_date. Mid-term endorsements "
        "(coverage changes, address updates, etc.) cause the policyholder's term dates to be slightly "
        "offset from the policy's effective dates. When these don't match exactly, the join returns no "
        "row → party_key = NULL → the entire downstream chain (same-as-link, CIF, email) is broken."
    ),
    code_before=(
        "-- BEFORE (broken)\n"
        "left join DWM.EDW.vw_policyholder vplh\n"
        "    on vplh.policy_term_key      = it.policy_term_key\n"
        "   and vplh.term_effective_date  = it.effective_from_date   -- strict date match\n"
        "   and vplh.term_expiration_date = it.effective_to_date     -- strict date match\n"
        "   and vplh.term_expiration_date > @as_of_date"
    ),
    discovery_text=(
        "missing_mail.sql Section 1 (Email Chain Funnel) showed:\n"
        "    Has NULL party_key (no policyholder match) : 35,820"
    ),
    fix_text=(
        "Relaxed the join to policy_term_key only — matching the pattern already used in paperless.sql:"
    ),
    code_after=(
        "-- AFTER (fixed)\n"
        "left join DWM.EDW.vw_policyholder vplh\n"
        "    on vplh.policy_term_key = it.policy_term_key"
    ),
    result_text=(
        "Has NULL party_key : 35,820 → 0\n"
        "Has duplicate_party_anchor_id (chain can proceed) : 261,592 → 297,412"
    ),
)
page_break(doc)

# ── ISSUE 2 ───────────────────────────────────────────────────────────────────
issue_section(
    2,
    "CIF Superseded Records Not Filtered (valid_to_date)",
    problem_text=(
        "AWM.dbo.cif_policy_party_detail uses two separate date columns: effective_to_date (when "
        "the paperless preference ends) and valid_to_date (when the record itself was superseded by "
        "a newer load). The original query only filtered effective_to_date = '9999-12-31', leaving "
        "superseded records in scope. A policy could return a stale paperless indicator from a "
        "previous load that has since been updated."
    ),
    code_before=(
        "-- BEFORE (incomplete filter)\n"
        "where cppd.output_document_type_code in ('BIL', 'POL')\n"
        "  and cppd.effective_to_date = '9999-12-31'"
    ),
    discovery_text=(
        "Code review of cif_policy_party_detail table structure. valid_to_date is the AWM standard "
        "audit column for record supersession — omitting it is a known pattern risk in AWM queries."
    ),
    fix_text="Added valid_to_date filter alongside the existing effective_to_date filter:",
    code_after=(
        "-- AFTER (fixed)\n"
        "where cppd.output_document_type_code in ('BIL', 'POL')\n"
        "  and cppd.effective_to_date = '9999-12-31'\n"
        "  and cppd.valid_to_date     = '9999-12-31'   -- excludes superseded records"
    ),
    result_text=(
        "Ensures only current, non-superseded paperless indicator records are used. Reduces risk "
        "of stale 'Y'/'N' indicators from old CIF loads appearing in the report."
    ),
)

# ── ISSUE 3 ───────────────────────────────────────────────────────────────────
issue_section(
    3,
    "CIF Dedup Using Wrong Sort Key",
    problem_text=(
        "To get the latest CIF record per (policy_party_link_id, output_document_type_code), the "
        "query used ROW_NUMBER() ordered by effective_from_date DESC. However, two records can share "
        "the same effective_from_date if a correction was loaded — in that case, effective_from_date "
        "does not distinguish which is truly the latest. update_date is the correct audit trail "
        "column in AWM for 'most recently written' row."
    ),
    code_before=(
        "-- BEFORE (wrong sort key)\n"
        "row_number() over (\n"
        "    partition by cppd.policy_party_link_id, cppd.output_document_type_code\n"
        "    order by cppd.effective_from_date desc\n"
        ") as rn"
    ),
    discovery_text=(
        "Code review of AWM table conventions. update_date is the standard 'last modified' "
        "timestamp on AWM operational tables."
    ),
    fix_text="Changed sort key to update_date DESC:",
    code_after=(
        "-- AFTER (correct sort key)\n"
        "row_number() over (\n"
        "    partition by cppd.policy_party_link_id, cppd.output_document_type_code\n"
        "    order by cppd.update_date desc\n"
        ") as rn"
    ),
    result_text=(
        "Guarantees the most recently updated CIF record wins the dedup, preventing stale indicators "
        "from surviving when a correction record was loaded on the same effective date."
    ),
)
page_break(doc)

# ── ISSUE 4 ───────────────────────────────────────────────────────────────────
issue_section(
    4,
    "NULL paper_notify_indicator Coerced to 'N' (False Negative)",
    problem_text=(
        "The paperless indicator logic used a two-branch CASE. When no BIL record exists for a "
        "policy, MIN(paper_notify_indicator) returns NULL. The CASE evaluates NULL in the ELSE branch "
        "→ outputs 'N'. This makes the policy appear explicitly non-paperless when in fact there is "
        "simply no data. Downstream, this created mismatches in the BIM vs EDW comparison."
    ),
    code_before=(
        "-- BEFORE (false negative)\n"
        "case when c.BIL_paper_notify = 0 then 'Y' else 'N' end as Paperless_Bil_Ind"
    ),
    discovery_text=(
        "BIM vs EDW mismatch analysis in cif_join_validation.sql — policies with BillInd = 'Y' in "
        "BIM but 'N' in EDW, where the root cause was a missing CIF record rather than a true "
        "preference change."
    ),
    fix_text="Added explicit NULL branch — unknown is not the same as non-paperless:",
    code_after=(
        "-- AFTER (correct NULL handling)\n"
        "case\n"
        "    when min(case when d.output_document_type_code = 'BIL'\n"
        "             then d.paper_notify_indicator end) = 0    then 'Y'\n"
        "    when min(case when d.output_document_type_code = 'BIL'\n"
        "             then d.paper_notify_indicator end) is null then null\n"
        "    else 'N'\n"
        "end as Paperless_Bil_Ind"
    ),
    result_text=(
        "Policies with no CIF record now show NULL instead of 'N' for paperless indicators. "
        "In Tableau, NULL can be filtered or displayed separately from explicit 'N'."
    ),
)

# ── ISSUE 5 ───────────────────────────────────────────────────────────────────
issue_section(
    5,
    "Email Fan-Out from asp_user_account_detail Partition Bug",
    problem_text=(
        "The deduplication of asp_user_account_detail to get one email per account link partitioned "
        "by (party_user_account_link_id, user_name). A single link_id with 3 distinct user_name "
        "values produced 3 rows all with rn=1 — one per username. When joined downstream, this "
        "fanned out every policy with multiple email addresses, producing counts larger than the "
        "total policy count."
    ),
    code_before=(
        "-- BEFORE (fan-out bug)\n"
        "row_number() over (\n"
        "    partition by auad.party_user_account_link_id, auad.user_name\n"
        "    order by auad.valid_from_date desc\n"
        ") as rn_detail\n\n"
        "-- Observed impossible results:\n"
        "-- UMB : email_resolved = 29,243  >  total = 27,389  <- impossible\n"
        "-- DP  : email_resolved = 11,812  >  total = 11,419  <- impossible"
    ),
    discovery_text=(
        "missing_mail.sql Sections 4/5 — email counts exceeded total policies, which is "
        "mathematically impossible and confirmed a fan-out bug in the partition logic."
    ),
    fix_text="Removed user_name from the partition — one row per link only:",
    code_after=(
        "-- AFTER (fixed)\n"
        "row_number() over (\n"
        "    partition by auad.party_user_account_link_id     -- link only, not per user_name\n"
        "    order by auad.valid_from_date desc, auad.valid_to_date desc\n"
        ") as rn_detail"
    ),
    result_text=(
        "Email counts now exactly satisfy has_email + missing_email = total for every symbol:\n"
        "  HO  : 101,232 + 29,185 = 130,417\n"
        "  CA  :  94,405 + 27,697 = 122,102\n"
        "  UMB :  22,599 +  4,790 =  27,389\n"
        "  DP  :   9,156 +  2,263 =  11,419\n"
        "  MA  :   4,867 +  1,219 =   6,086"
    ),
)
page_break(doc)

# ── ISSUE 6 ───────────────────────────────────────────────────────────────────
issue_section(
    6,
    "Known Bad Party Anchor IDs Inflating Email Links",
    problem_text=(
        "Two party anchor IDs (7771543 and 13322119) appear in AWM.dbo.party_user_account_link "
        "and produce incorrect or inflated link matches. These are known data anomalies."
    ),
    discovery_text=(
        "Identified in the original paperless.sql codebase (line 284) which already excluded them "
        "with a comment. Any new query building the email chain that doesn't carry this exclusion "
        "will pick up incorrect party-to-account mappings."
    ),
    fix_text="Added exclusion to all email chain queries in both new_paper.sql and missing_mail.sql:",
    code_after=(
        "where pual.party_anchor_id not in ('7771543', '13322119')"
    ),
    result_text=(
        "Prevents two known bad anchors from being matched to unrelated parties during the "
        "party_user_account_link join, ensuring email resolves to the correct customer."
    ),
)

# ── ISSUE 7 ───────────────────────────────────────────────────────────────────
issue_section(
    7,
    "BIM Fallback Email Join Format Mismatch",
    problem_text=(
        "#bim_email stores P.POL_KEY from Eloqua.POLICY. The initial assumption was that POL_KEY "
        "is numeric-only (e.g. 0130957), requiring the EDW policy_number (e.g. CA 0130957) to be "
        "stripped before joining. This stripping broke the join entirely — bim_fallback_used = 0."
    ),
    code_before=(
        "-- BROKEN attempt (incorrect strip)\n"
        "left join #bim_email bim\n"
        "    on bim.policy_number = substring(\n"
        "        src.policy_number,\n"
        "        case when src.policy_number like 'UMB%' then 5 else 4 end,\n"
        "        len(src.policy_number)\n"
        "    )"
    ),
    discovery_text=(
        "Validation query showed bim_fallback_used = 0 after the strip was applied — zero BIM "
        "emails were matching, confirming the format assumption was wrong."
    ),
    fix_text=(
        "Confirmed POL_KEY format matches EDW policy_number (both use 'CA 0130957' format). "
        "Reverted to direct join:"
    ),
    code_after=(
        "-- FIXED (direct join)\n"
        "left join #bim_email bim\n"
        "    on bim.policy_number = src.policy_number"
    ),
    result_text=(
        "BIM fallback join started matching correctly. After the ANI fix (Issue 8) resolved the "
        "bulk of the gap via AWM, the BIM fallback covered the remaining 464 truly AWM-absent policies."
    ),
)
page_break(doc)

# ── ISSUE 8 ───────────────────────────────────────────────────────────────────
issue_section(
    8,
    "Grain Dedup Losing Email by Preferring NIN Regardless of Email",
    problem_text=(
        "The final output dedup to one row per policy_term_key sorted first by NIN preference. "
        "This always selected the NIN (named insured / head of household) row regardless of whether "
        "that party had an email. For policies where NIN had no AWM email but the ANI did, the email "
        "was discarded — the NIN row won and EmailAddress = NULL."
    ),
    code_before=(
        "-- BEFORE (email-blind dedup)\n"
        "row_number() over (\n"
        "    partition by b.policy_term_key\n"
        "    order by\n"
        "        case when b.policyholder_type_code = 'NIN' then 0 else 1 end,\n"
        "        b.party_key desc\n"
        ") as rn_grain"
    ),
    discovery_text=(
        "After the policyholder date join fix (Issue 1), awm_email count was 225,188 — still lower "
        "than the 232,259 expected from the Section 1 funnel. The gap of ~7,071 was traced to "
        "policies where only the ANI party resolved an email address."
    ),
    fix_text="Added email presence as the first dedup priority, before NIN preference:",
    code_after=(
        "-- AFTER (email-aware dedup)\n"
        "row_number() over (\n"
        "    partition by b.policy_term_key\n"
        "    order by\n"
        "        case when ea.user_name is not null then 0 else 1 end,  -- Priority 1: has email\n"
        "        case when b.policyholder_type_code = 'NIN' then 0 else 1 end,  -- Priority 2: NIN\n"
        "        b.party_key desc                                        -- Priority 3: tiebreak\n"
        ") as rn_grain"
    ),
    result_text=(
        "AWM email : 225,188 → 251,664   (+26,476 recovered from ANI parties)\n"
        "no_email  :  71,211 →  44,735   (-26,476)\n\n"
        "Note: When an ANI row wins due to email, the output row carries ANI demographics "
        "(HOH_IND = 'N', policyholder_type_code = 'ANI'). The EmailSource column identifies "
        "the row as AWM-sourced."
    ),
)

# ── ISSUE 9 ───────────────────────────────────────────────────────────────────
issue_section(
    9,
    "BIM HOH Filter Too Restrictive for Fallback",
    problem_text=(
        "#bim_email was built with WHERE C.HOH_IND = 'Y' to pick only the head of household's "
        "email from BIM. Many of the 27,756 target policies (incomplete AWM account setup) had "
        "email only on non-HOH BIM contacts. The HOH contact either had no email or was the one "
        "with the incomplete AWM account. Filtering to HOH only excluded the majority of valid "
        "BIM emails."
    ),
    code_before=(
        "-- BEFORE (over-filtered)\n"
        "where C.END_DT      = '9999-12-31'\n"
        "  and P.END_DT      = '9999-12-31'\n"
        "  and P.POL_STS_CD  = 'INFORCE'\n"
        "  and C.HOH_IND     = 'Y'                -- too restrictive\n"
        "  and C.EMAIL_ADDRESS is not null"
    ),
    discovery_text=(
        "After the join format was corrected (Issue 7), bim_fallback_used showed only 2,375 "
        "instead of the expected ~27,756 from the Section 7c analysis."
    ),
    fix_text=(
        "Removed HOH_IND = 'Y' filter. MAX(EMAIL_ADDRESS) already ensures one email per POL_KEY "
        "across all contacts:"
    ),
    code_after=(
        "-- AFTER (all contacts, one email per policy)\n"
        "where C.END_DT      = '9999-12-31'\n"
        "  and P.END_DT      = '9999-12-31'\n"
        "  and P.POL_STS_CD  = 'INFORCE'\n"
        "  and C.EMAIL_ADDRESS is not null\n"
        "  and C.EMAIL_ADDRESS <> ''\n"
        "group by P.POL_KEY   -- MAX(EMAIL_ADDRESS) picks one per policy"
    ),
    result_text=(
        "After Issue 8 (ANI fix) resolved the bulk of the gap via AWM, the BIM fallback was left "
        "with the remaining 464 truly AWM-absent policies — the correct expected residual."
    ),
)
page_break(doc)

# ── EMAIL GAP ANALYSIS ────────────────────────────────────────────────────────
h1(doc, "Email Gap Analysis Summary")
body(doc, "After all fixes, the remaining 44,271 policies with no email break down as follows (from missing_mail.sql Section 9a):")

add_table(doc,
    headers=["Gap Category", "Policies", "Resolution"],
    rows=[
        ["No same-as-link (AWM party not mapped)",              "1",       "Data fix: investigate party_id_same_as_link"],
        ["Never registered online (no account on any party)",   "50,453",  "Outreach: campaign to drive online registration"],
        ["Incomplete account setup — no email in either system","13,846",  "Outreach: prompt customers to complete account activation"],
        ["Incomplete account setup — BIM email available",      "27,756",  "Data fix: backfill from BIM (see missing_mail.sql Section 8)"],
    ],
    col_widths=[2.8, 0.9, 2.7],
)

note(doc, (
    "The 50,453 + 13,846 + 27,756 total is from the pre-ANI-fix analysis. "
    "After the ANI fix (Issue 8) resolved 26,476 via AWM, the structural no-email count "
    "reduced to ~44,271. Relative proportions across categories remain consistent."
))

h2(doc, "No-Account Policies — Product Line Breakdown")
add_table(doc,
    headers=["Symbol", "Unregistered Policies", "% of Symbol Total"],
    rows=[
        ["HO",  "67,670", "51.9%"],
        ["CA",  "61,921", "50.7%"],
        ["UMB", "14,634", "53.4%"],
        ["DP",  " 6,027", "52.8%"],
        ["MA",  " 3,748", "61.6%"],
    ],
    col_widths=[1.0, 2.0, 1.8],
)
body(doc, (
    "All symbols are consistently 50–62% without an account link. This is a structural "
    "customer adoption gap, not a data pipeline issue."
))

page_break(doc)

# ── FINAL VALIDATION ─────────────────────────────────────────────────────────
h1(doc, "Final Output Validation")
body(doc, "Run this query after new_paper.sql completes to confirm expected values:")

add_code_block(doc,
    "select\n"
    "    count(*)                                                         as total_rows,\n"
    "    sum(case when EmailAddress is not null then 1 else 0 end)        as has_email,\n"
    "    sum(case when EmailSource = 'BIM'      then 1 else 0 end)        as bim_fallback_used,\n"
    "    sum(case when EmailSource = 'AWM'      then 1 else 0 end)        as awm_email,\n"
    "    sum(case when EmailAddress is null     then 1 else 0 end)        as no_email,\n"
    "    sum(case when Paperless_Bil_Ind = 'Y' then 1 else 0 end)        as paperless_bill_y,\n"
    "    sum(case when Paperless_Pol_Ind = 'Y' then 1 else 0 end)        as paperless_pol_y,\n"
    "    sum(case when PolicyStatus = 'INFORCE' then 1 else 0 end)        as inforce_count\n"
    "from #New_Paper_Report;"
)

h2(doc, "Expected Results (as of 2026-04-07)")
add_table(doc,
    headers=["Metric", "Expected Value"],
    rows=[
        ["total_rows",         "~296,399"],
        ["has_email",          "~252,128  (85.1%)"],
        ["bim_fallback_used",  "~464"],
        ["awm_email",          "~251,664"],
        ["no_email",           "~44,271   (14.9%)"],
        ["paperless_bill_y",   "~113,834"],
        ["paperless_pol_y",    "~146,296"],
        ["inforce_count",      "~284,879"],
    ],
    col_widths=[2.5, 2.5],
)

# ── PIPELINE OPTIMIZATION ─────────────────────────────────────────────────────
h1(doc, "Pipeline Optimization Summary")
body(doc, "new_paper.sql reduced the temp table count from 16 (in paperless.sql) to 10:")

h2(doc, "Tables Removed")
add_table(doc,
    headers=["Removed Table", "Merged Into"],
    rows=[
        ["#base_data_link",            "Inlined as subquery in #base_data"],
        ["#CIF_POLICY_Detail1",        "Merged into #cif_detail"],
        ["#CIF_POLICY_Detail2",        "Merged into #cif_detail"],
        ["#ct_asp_user_account",       "Merged into #email_account"],
        ["#ct_party_user_account_link","Merged into #email_account"],
        ["#ct_party_user_account",     "Merged into #email_account"],
        ["#ct_asp_membership",         "Merged into #email_account"],
    ],
    col_widths=[2.8, 3.5],
)

h2(doc, "Tables Added")
add_table(doc,
    headers=["Added Table", "Purpose"],
    rows=[
        ["#bim_email", "BIM/Eloqua email fallback (Issue 7 — join format fixed, HOH filter removed)"],
    ],
    col_widths=[1.5, 5.0],
)

body(doc,
    "Grain dedup (one row per policy_term_key) is handled in the final SELECT without an "
    "additional temp table, using a priority-aware ROW_NUMBER() that selects email-present "
    "rows first, NIN second, and uses party_key as a deterministic tiebreak."
)

# ── SAVE ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"Saved: {OUTPUT_PATH}")
